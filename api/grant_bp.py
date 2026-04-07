"""
ZEREK — Генератор бизнес-плана на грант 400 МРП (Бастау Бизнес)
Заполняет шаблон .docx данными из анкеты + расчётами.
"""
import os
from io import BytesIO
from docx import Document

MRP_2026 = 4325
GRANT_400MRP = 400 * MRP_2026  # 1 730 000 ₸

# ═══════════════════════════════════════════════
# СПРАВОЧНИКИ
# ═══════════════════════════════════════════════

CITY_DATA = {
    "astana":    {"name": "Астана",    "region": "город Астана",              "avg_wage": 400000},
    "almaty":    {"name": "Алматы",    "region": "город Алматы",              "avg_wage": 420000},
    "shymkent":  {"name": "Шымкент",   "region": "город Шымкент",             "avg_wage": 270000},
    "aktobe":    {"name": "Актобе",    "region": "Актюбинская область",       "avg_wage": 320000},
    "karaganda": {"name": "Караганда", "region": "Карагандинская область",    "avg_wage": 330000},
    "atyrau":    {"name": "Атырау",    "region": "Атырауская область",       "avg_wage": 380000},
    "pavlodar":  {"name": "Павлодар",  "region": "Павлодарская область",     "avg_wage": 310000},
    "kostanay":  {"name": "Костанай",  "region": "Костанайская область",     "avg_wage": 290000},
    "aktau":     {"name": "Актау",     "region": "Мангистауская область",    "avg_wage": 370000},
    "semey":     {"name": "Семей",     "region": "область Абай",             "avg_wage": 280000},
    "uralsk":    {"name": "Уральск",   "region": "Западно-Казахстанская область", "avg_wage": 300000},
    "taldykorgan": {"name": "Талдыкорган", "region": "Алматинская область",  "avg_wage": 275000},
    "turkestan": {"name": "Туркестан", "region": "Туркестанская область",    "avg_wage": 250000},
    "kokshetau": {"name": "Кокшетау",  "region": "Акмолинская область",     "avg_wage": 285000},
    "petropavl": {"name": "Петропавловск", "region": "Северо-Казахстанская область", "avg_wage": 290000},
}

FORMAT_NAMES = {
    "COFFEE_KIOSK": "кофе-точка (кофе с собой)",
    "COFFEE_ISLAND": "кофе-остров (островок в ТЦ)",
    "COFFEE_MINI": "мини-кофейня",
    "COFFEE_FULL": "кофейня-кафе",
    "COFFEE_VENDING": "кофе-вендинг",
    "COFFEE_MOBILE": "кофе на колёсах",
    "DONER_KIOSK": "донерная (точка на вынос)",
    "DONER_MINI": "мини-кафе (донерная с залом)",
    "DONER_CAFE": "донер-кафе",
    "DONER_DELIVERY": "доставка (дарк китчен)",
    "CARWASH_MANUAL_S": "автомойка (1-2 бокса)",
    "CARWASH_MANUAL_L": "автомойка (3-5 боксов)",
    "CARWASH_SELF_BOX": "мойка самообслуживания (бокс)",
    "BARBER_SOLO": "барбершоп (аренда кресла)",
    "BARBER_MINI": "мини-барбершоп (2 кресла)",
    "BARBER_FULL": "барбершоп (3+ кресла)",
    "CLEAN_SOLO": "клининг (самозанятый)",
    "CLEAN_TEAM": "клининг (бригада)",
    "GROCERY_MINI": "мини-маркет",
    "GROCERY_SPEC": "специализированный магазин",
    "PHARMA_MINI": "аптечный пункт",
    "PHARMA_STD": "аптека",
    "NAIL_HOME": "маникюр (на дому)",
    "NAIL_CABINET": "маникюрный кабинет",
    "NAIL_STUDIO": "маникюрная студия",
    "BAKERY_MINI": "мини-пекарня",
    "BAKERY_ISLAND": "пекарня-островок в ТЦ",
    "PIZZA_DELIVERY": "доставка пиццы",
    "PIZZA_TAKEAWAY": "пиццерия на вынос",
    "DENTAL_CABINET": "стоматологический кабинет",
    "TIRE_MINI": "мини-шиномонтаж",
    "TIRE_MOBILE": "мобильный шиномонтаж",
    "BROW_HOME": "брови (на дому)",
    "BROW_CABINET": "бровный кабинет",
    "LASH_HOME": "ресницы (на дому)",
    "LASH_CABINET": "кабинет ресниц",
    "SUGARING_HOME": "шугаринг (на дому)",
    "SUGARING_OWN": "свой кабинет шугаринга",
    "MASSAGE_CABINET": "массажный кабинет",
    "FITNESS_STUDIO": "фитнес-студия",
    "FLOWERS_KIOSK": "цветочный ларёк",
    "FLOWERS_SHOP": "цветочный магазин",
    "FV_KIOSK": "овощной киоск",
    "FURN_SOLO": "мебель (мастер-одиночка)",
    "PVZ_SMALL": "мини-ПВЗ",
    "REPAIR_PHONE_KIOSK": "ремонт телефонов (киоск)",
    "REP_KIOSK": "ремонт телефонов (точка в ТЦ)",
    "REP_SHOP": "сервисный центр",
    "KINDER_HOME": "домашний детский сад",
    "CYBER_MINI": "мини-компьютерный клуб",
    "DRY_POINT": "приёмный пункт химчистки",
    "DRY_MINI": "мини-химчистка",
    "TAILOR_HOME": "ателье (на дому)",
    "TAILOR_MINI": "мини-ателье",
    "WATER_VENDING": "водомат",
    "WATER_FILTER": "фильтрация и розлив воды",
    "SEMI_HOME": "полуфабрикаты (домашнее производство)",
    "SEMI_MINI": "мини-цех полуфабрикатов",
    "CONF_HOME": "домашний кондитер",
    "CONF_CABINET": "кондитерский кабинет",
    "CANTEEN_OFFICE": "столовая при БЦ",
    "CANTEEN_CITY": "городская столовая",
    "FASTFOOD_KIOSK": "фастфуд-киоск",
    "FASTFOOD_MINI": "мини-кафе быстрого питания",
    "SUSHI_DELIVERY": "доставка суши",
    "SUSHI_BAR": "суши-бар",
}

# ═══════════════════════════════════════════════
# ДАННЫЕ ПО НИШАМ
# ═══════════════════════════════════════════════

NICHE_DATA = {
    "COFFEE": {
        "name_ru": "Кофейня",
        "goal_template": "Открытие кофейни формата «{format_name}» в городе {city_name} для обеспечения жителей качественными кофейными напитками и выпечкой",
        "description_template": "Проект предусматривает открытие кофейни формата «{format_name}». Планируется реализация кофейных напитков (американо, капучино, латте), чая, выпечки и десертов. Целевая аудитория — жители района и проходящий трафик.",
        "services": [
            {"name": "Кофейные напитки (американо, латте, капучино)", "price": 800, "daily_volume": 40},
            {"name": "Чай, какао", "price": 500, "daily_volume": 10},
            {"name": "Выпечка, десерты", "price": 600, "daily_volume": 20},
        ],
        "capex_items": {
            "COFFEE_KIOSK": [
                {"name": "Кофемашина профессиональная (2 группы)", "qty": 1, "price": 650000},
                {"name": "Кофемолка", "qty": 1, "price": 120000},
                {"name": "Холодильная витрина", "qty": 1, "price": 180000},
                {"name": "Барная стойка", "qty": 1, "price": 150000},
                {"name": "Посуда, инвентарь", "qty": 1, "price": 80000},
                {"name": "Кассовый аппарат + POS-терминал", "qty": 1, "price": 85000},
                {"name": "Вывеска, оформление", "qty": 1, "price": 120000},
                {"name": "Первоначальный запас сырья (зёрна, молоко, сиропы)", "qty": 1, "price": 150000},
            ],
            "COFFEE_ISLAND": [
                {"name": "Кофемашина профессиональная (2 группы)", "qty": 1, "price": 650000},
                {"name": "Кофемолка", "qty": 1, "price": 120000},
                {"name": "Островной модуль (стойка + витрина)", "qty": 1, "price": 350000},
                {"name": "Посуда, инвентарь", "qty": 1, "price": 80000},
                {"name": "Кассовый аппарат + POS-терминал", "qty": 1, "price": 85000},
                {"name": "Вывеска, оформление", "qty": 1, "price": 100000},
                {"name": "Первоначальный запас сырья", "qty": 1, "price": 150000},
            ],
        },
        "opex_monthly": {
            "COFFEE_KIOSK": {"rent": 150000, "fot": 170000, "raw_materials": 180000, "utilities": 25000, "other": 30000},
            "COFFEE_ISLAND": {"rent": 200000, "fot": 170000, "raw_materials": 200000, "utilities": 15000, "other": 30000},
        },
        "competitors_text": "В городе работают сетевые кофейни (Coffeedelia, Illy Coffee), а также локальные точки. Конкуренция средняя.",
        "advantage_text": "Свежая обжарка, стабильное качество напитков, удобная локация с высоким трафиком, доступные цены.",
        "suppliers_text": "Поставщики: Paulig Kazakhstan (кофе в зёрнах), местные пекарни (выпечка), FoodMaster/Lactel (молочная продукция).",
        "location_reasons": {
            "ТЦ": "Высокий пешеходный трафик, наличие парковки, целевая аудитория — посетители ТЦ.",
            "жилой_район": "Стабильный поток жителей микрорайона, утренний и вечерний трафик.",
            "бизнес_центр": "Высокая концентрация офисных работников — целевая аудитория для кофе.",
            "рынок": "Высокий трафик покупателей, особенно в выходные дни.",
            "улица": "Проходная улица с высоким пешеходным и автомобильным трафиком.",
        },
        "channels_text": "Instagram, 2ГИС, Яндекс Карты, вывеска на точке, сарафанное радио.",
        "consumer_annual_purchases": "В среднем 1 клиент покупает кофе 3–4 раза в неделю, ~150 раз в год, средний чек 900 ₸, итого ~135 000 ₸ в год.",
        "data_sources": "Данные Бюро национальной статистики РК, анализ 2ГИС, собственные наблюдения.",
    },
    "DONER": {
        "name_ru": "Донерная",
        "goal_template": "Открытие точки быстрого питания (донерная) формата «{format_name}» в городе {city_name}",
        "description_template": "Проект предусматривает открытие точки быстрого питания, специализирующейся на приготовлении донеров, шаурмы и сопутствующих блюд. Формат — «{format_name}». Целевая аудитория — жители района, студенты, работающее население.",
        "services": [
            {"name": "Донер/шаурма", "price": 1200, "daily_volume": 50},
            {"name": "Комбо-набор (донер + напиток)", "price": 1600, "daily_volume": 15},
            {"name": "Напитки (чай, лимонад, вода)", "price": 300, "daily_volume": 30},
        ],
        "capex_items": {
            "DONER_KIOSK": [
                {"name": "Донер-аппарат (шаурма-гриль)", "qty": 1, "price": 250000},
                {"name": "Холодильник", "qty": 1, "price": 180000},
                {"name": "Фритюрница", "qty": 1, "price": 65000},
                {"name": "Рабочий стол из нержавейки", "qty": 1, "price": 85000},
                {"name": "Мойка", "qty": 1, "price": 45000},
                {"name": "Кассовый аппарат + POS-терминал", "qty": 1, "price": 85000},
                {"name": "Вывеска, оформление", "qty": 1, "price": 100000},
                {"name": "Посуда, инвентарь, упаковка", "qty": 1, "price": 70000},
                {"name": "Первоначальный запас продуктов", "qty": 1, "price": 200000},
            ],
            "DONER_MINI": [
                {"name": "Донер-аппарат (шаурма-гриль)", "qty": 1, "price": 250000},
                {"name": "Холодильник промышленный", "qty": 1, "price": 250000},
                {"name": "Фритюрница", "qty": 1, "price": 65000},
                {"name": "Плита газовая 2-конфорочная", "qty": 1, "price": 75000},
                {"name": "Рабочие столы из нержавейки", "qty": 2, "price": 85000},
                {"name": "Мойка 2-секционная", "qty": 1, "price": 65000},
                {"name": "Кассовый аппарат + POS-терминал", "qty": 1, "price": 85000},
                {"name": "Вывеска, оформление", "qty": 1, "price": 120000},
                {"name": "Мебель для зала (столы, стулья)", "qty": 1, "price": 200000},
                {"name": "Посуда, инвентарь, упаковка", "qty": 1, "price": 80000},
                {"name": "Первоначальный запас продуктов", "qty": 1, "price": 250000},
            ],
        },
        "opex_monthly": {
            "DONER_KIOSK": {"rent": 120000, "fot": 200000, "raw_materials": 350000, "utilities": 30000, "other": 30000},
            "DONER_MINI": {"rent": 180000, "fot": 350000, "raw_materials": 500000, "utilities": 40000, "other": 40000},
        },
        "competitors_text": "Рынок быстрого питания насыщен — сети Doner King, Sultan Doner, множество локальных точек. Конкуренция высокая.",
        "advantage_text": "Свежее мясо от местных поставщиков, большие порции, быстрое обслуживание (до 5 минут), доступные цены.",
        "suppliers_text": "Поставщики: местные мясокомбинаты (курица, говядина), овощные базы, оптовые рынки.",
        "location_reasons": {
            "ТЦ": "Высокий трафик, фуд-корт, импульсный спрос.",
            "жилой_район": "Стабильный вечерний спрос, доставка по району.",
            "рынок": "Максимальный трафик, особенно в обеденное время и выходные.",
            "улица": "Проходная улица рядом с остановкой/учебными заведениями.",
            "бизнес_центр": "Обеденный трафик офисных работников.",
        },
        "channels_text": "Instagram, 2ГИС, Яндекс Карты, Glovo/Wolt, вывеска, листовки.",
        "consumer_annual_purchases": "Средний клиент покупает донер 2 раза в неделю, ~100 раз в год, средний чек 1 400 ₸, итого ~140 000 ₸ в год.",
        "data_sources": "Данные БНС РК по общепиту, анализ 2ГИС, мониторинг цен конкурентов.",
    },
    "BARBER": {
        "name_ru": "Барбершоп",
        "goal_template": "Открытие барбершопа формата «{format_name}» в городе {city_name} для оказания услуг мужских стрижек и ухода",
        "description_template": "Проект предусматривает открытие барбершопа формата «{format_name}». Услуги: мужские стрижки, бритьё, моделирование бороды, детские стрижки. Целевая аудитория — мужчины 18-50 лет.",
        "services": [
            {"name": "Мужская стрижка", "price": 3000, "daily_volume": 8},
            {"name": "Стрижка + борода", "price": 4500, "daily_volume": 3},
            {"name": "Детская стрижка", "price": 2000, "daily_volume": 2},
        ],
        "capex_items": {
            "BARBER_MINI": [
                {"name": "Парикмахерское кресло", "qty": 2, "price": 120000},
                {"name": "Зеркало с подсветкой", "qty": 2, "price": 45000},
                {"name": "Машинки для стрижки (набор)", "qty": 2, "price": 80000},
                {"name": "Мойка для головы", "qty": 1, "price": 150000},
                {"name": "Инструменты (ножницы, бритвы, расчёски)", "qty": 1, "price": 60000},
                {"name": "Стерилизатор", "qty": 1, "price": 35000},
                {"name": "Мебель (ресепшн, зона ожидания)", "qty": 1, "price": 150000},
                {"name": "Кассовый аппарат + POS-терминал", "qty": 1, "price": 85000},
                {"name": "Вывеска, оформление", "qty": 1, "price": 120000},
                {"name": "Первоначальный запас расходников", "qty": 1, "price": 80000},
            ],
        },
        "opex_monthly": {
            "BARBER_MINI": {"rent": 150000, "fot": 300000, "raw_materials": 40000, "utilities": 25000, "other": 25000},
        },
        "competitors_text": "В городе работают сети TopGun, OldBoy и локальные барбершопы. Конкуренция растёт, но спрос стабильный.",
        "advantage_text": "Качественные стрижки по доступным ценам, удобное расположение, система записи через Instagram/WhatsApp.",
        "suppliers_text": "Поставщики: Wahl Kazakhstan (машинки), Moser (инструменты), локальные поставщики косметики для барберов.",
        "location_reasons": {
            "ТЦ": "Высокий трафик, удобная парковка.",
            "жилой_район": "Шаговая доступность для жителей, постоянные клиенты.",
            "бизнес_центр": "Обеденный перерыв — удобное время для стрижки.",
            "рынок": "Высокий трафик, доступная аренда.",
            "улица": "Проходная локация, вывеска привлекает проходящих.",
        },
        "channels_text": "Instagram, 2ГИС, WhatsApp-запись, вывеска, сарафанное радио.",
        "consumer_annual_purchases": "Мужчина стрижётся в среднем раз в 3-4 недели, ~12 раз в год. Средний чек 3 500 ₸, итого ~42 000 ₸ в год.",
        "data_sources": "Данные БНС РК, анализ 2ГИС, мониторинг цен конкурентов.",
    },
    "CARWASH": {
        "name_ru": "Автомойка",
        "goal_template": "Открытие автомойки формата «{format_name}» в городе {city_name}",
        "description_template": "Проект предусматривает открытие автомойки формата «{format_name}». Услуги: мойка кузова, салона, полировка, химчистка. Целевая аудитория — автовладельцы города.",
        "services": [
            {"name": "Мойка кузова (стандарт)", "price": 2500, "daily_volume": 12},
            {"name": "Комплексная мойка (кузов + салон)", "price": 5000, "daily_volume": 4},
            {"name": "Химчистка салона", "price": 15000, "daily_volume": 1},
        ],
        "capex_items": {
            "CARWASH_MANUAL_S": [
                {"name": "Аппарат высокого давления (Karcher)", "qty": 1, "price": 350000},
                {"name": "Пылесос промышленный", "qty": 1, "price": 120000},
                {"name": "Пеногенератор", "qty": 1, "price": 85000},
                {"name": "Компрессор", "qty": 1, "price": 150000},
                {"name": "Химия и расходники (первоначальный запас)", "qty": 1, "price": 100000},
                {"name": "Освещение бокса", "qty": 1, "price": 60000},
                {"name": "Водоотвод, решётки", "qty": 1, "price": 120000},
                {"name": "Вывеска", "qty": 1, "price": 80000},
                {"name": "Инвентарь (губки, тряпки, вёдра)", "qty": 1, "price": 40000},
            ],
        },
        "opex_monthly": {
            "CARWASH_MANUAL_S": {"rent": 120000, "fot": 250000, "raw_materials": 60000, "utilities": 40000, "other": 20000},
        },
        "competitors_text": "Автомойки — распространённая ниша. В каждом районе 3-5 точек. Конкуренция высокая, но спрос стабильный.",
        "advantage_text": "Качественная химия, бережное отношение к автомобилю, быстрое обслуживание (30-40 минут), удобная локация.",
        "suppliers_text": "Поставщики: Karcher Kazakhstan (оборудование), Koch Chemie (автохимия), местные оптовые базы.",
        "location_reasons": {
            "жилой_район": "Стабильный поток автовладельцев микрорайона.",
            "улица": "Высокий автомобильный трафик, удобный заезд.",
            "рынок": "Рядом с рынком — высокий поток автомобилей.",
            "ТЦ": "Парковка ТЦ — удобно для клиентов.",
            "бизнес_центр": "Автовладельцы моют машину во время работы.",
        },
        "channels_text": "2ГИС, Яндекс Карты, вывеска, Instagram, листовки на парковках.",
        "consumer_annual_purchases": "Автовладелец моет машину 2-3 раза в месяц, ~30 раз в год. Средний чек 3 000 ₸, итого ~90 000 ₸ в год.",
        "data_sources": "Данные комитета по статистике, анализ 2ГИС по количеству автомоек.",
    },
}

# Фоллбэк для ниш без детальных данных
DEFAULT_NICHE = {
    "name_ru": "Бизнес-проект",
    "goal_template": "Открытие бизнеса формата «{format_name}» в городе {city_name}",
    "description_template": "Проект предусматривает открытие бизнеса формата «{format_name}» в городе {city_name}. Целевая аудитория — жители города и района.",
    "services": [
        {"name": "Основная услуга/товар", "price": 2000, "daily_volume": 20},
        {"name": "Дополнительная услуга", "price": 1000, "daily_volume": 10},
    ],
    "capex_items": {},
    "opex_monthly": {},
    "competitors_text": "На рынке присутствуют как сетевые, так и локальные игроки. Конкуренция умеренная.",
    "advantage_text": "Качественный сервис, удобная локация, доступные цены, индивидуальный подход.",
    "suppliers_text": "Работа с проверенными поставщиками по договору.",
    "location_reasons": {"ТЦ": "Высокий трафик.", "жилой_район": "Стабильный поток клиентов.", "улица": "Проходная локация.", "рынок": "Высокий трафик.", "бизнес_центр": "Офисные работники."},
    "channels_text": "Instagram, 2ГИС, Яндекс Карты, вывеска, сарафанное радио.",
    "consumer_annual_purchases": "Средний клиент пользуется услугой регулярно.",
    "data_sources": "Данные БНС РК, анализ 2ГИС.",
}


# ═══════════════════════════════════════════════
# ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ
# ═══════════════════════════════════════════════

def fmt(n):
    """Форматирует число с пробелами: 1730000 → '1 730 000'."""
    if n is None:
        return "—"
    return f"{int(round(n)):,}".replace(",", " ")


def set_cell(table, row, col, text):
    """Заполняет ячейку таблицы, сохраняя форматирование."""
    try:
        cell = table.rows[row].cells[col]
    except (IndexError, KeyError):
        # Merged cells — try accessing via XML
        try:
            tr = table.rows[row]._tr
            tcs = tr.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tc')
            if col < len(tcs):
                from docx.table import _Cell
                cell = _Cell(tcs[col], table)
            else:
                return
        except Exception:
            return
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ""
    if cell.paragraphs:
        cell.paragraphs[0].text = str(text)
    else:
        cell.add_paragraph(str(text))


def replace_in_paragraph(paragraph, old_text, new_text):
    """Заменяет текст в параграфе (может быть разбит по runs)."""
    full = "".join(r.text for r in paragraph.runs)
    if old_text in full:
        new_full = full.replace(old_text, new_text)
        for i, r in enumerate(paragraph.runs):
            r.text = new_full if i == 0 else ""
        return True
    if old_text in paragraph.text:
        paragraph.text = paragraph.text.replace(old_text, new_text)
        return True
    return False


def split_grant_own(capex_items, grant_max, own_funds=0):
    """Разбивает CAPEX на грантовые и собственные средства."""
    grant_remaining = grant_max
    result = []
    for item in capex_items:
        item_total = item["qty"] * item["price"]
        grant_part = min(item_total, grant_remaining)
        own_part = item_total - grant_part
        grant_remaining -= grant_part
        result.append({
            **item,
            "total": item_total,
            "grant": grant_part,
            "own": own_part,
        })
    return result


def calc_profit_forecast(monthly_revenue, opex, start_month):
    """Прогноз прибыли на 3 года."""
    months_y1 = 13 - start_month
    total_opex = sum(opex.values())

    rev_y1 = monthly_revenue * months_y1
    exp_y1 = total_opex * months_y1

    rev_y2 = monthly_revenue * 12 * 1.05
    exp_y2 = total_opex * 12 * 1.03

    rev_y3 = monthly_revenue * 12 * 1.10
    exp_y3 = total_opex * 12 * 1.06

    return {
        "y1": {"revenue": int(rev_y1), "expense": int(exp_y1), "profit": int(rev_y1 - exp_y1)},
        "y2": {"revenue": int(rev_y2), "expense": int(exp_y2), "profit": int(rev_y2 - exp_y2)},
        "y3": {"revenue": int(rev_y3), "expense": int(exp_y3), "profit": int(rev_y3 - exp_y3)},
    }


# ═══════════════════════════════════════════════
# ГЕНЕРАТОР
# ═══════════════════════════════════════════════

def generate_grant_bp(
    template_path: str,
    fio: str,
    iin: str,
    phone: str,
    address: str,
    legal_status: str,
    legal_address: str,
    experience_years: int,
    family_status: str,
    city_id: str,
    niche_id: str,
    format_id: str,
    project_name: str,
    location_description: str,
    loc_type: str = "ТЦ",
    own_funds: int = 0,
    grant_amount: int = GRANT_400MRP,
    start_month: int = 1,
    **kwargs,
) -> bytes:
    """Генерирует заполненный бизнес-план .docx и возвращает bytes."""

    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Шаблон не найден: {template_path}")

    doc = Document(template_path)
    city = CITY_DATA.get(city_id, {"name": city_id, "region": city_id, "avg_wage": 300000})
    niche = NICHE_DATA.get(niche_id, DEFAULT_NICHE)
    format_name = FORMAT_NAMES.get(format_id, format_id)

    # ── Параграфы (заголовки) ──
    paragraphs = doc.paragraphs
    if len(paragraphs) > 0:
        for p in paragraphs:
            replace_in_paragraph(p, "Костанайская область/", f"{city['region']}/")
            replace_in_paragraph(p, "Костанайская область", city["region"])
            replace_in_paragraph(p, "город Лисаковск", f"город {city['name']}")
            replace_in_paragraph(p, "«Наименование бизнес-проект»", f"«{project_name}»")
            replace_in_paragraph(p, "Наименование бизнес-проект", project_name)
            replace_in_paragraph(p, "_____ фамилия, имя, отчество", fio)
            replace_in_paragraph(p, "фамилия, имя, отчество", fio)

    tables = doc.tables

    # ── Таблица 0: Инициатор проекта ──
    if len(tables) > 0:
        t = tables[0]
        personal = [
            fio,
            iin,
            phone,
            address,
            legal_status,
            legal_address,
            f"{experience_years} лет" if experience_years > 0 else "Без опыта",
            family_status,
        ]
        for i, val in enumerate(personal):
            if i < len(t.rows):
                set_cell(t, i, 1, val)

    # ── Таблица 1: Цель проекта ──
    if len(tables) > 1:
        goal = niche["goal_template"].format(
            format_name=format_name, city_name=city["name"]
        )
        set_cell(tables[1], 0, 0, goal)

    # ── Таблица 2: Краткая информация ──
    if len(tables) > 2:
        desc = niche["description_template"].format(
            format_name=format_name, city_name=city["name"]
        )
        set_cell(tables[2], 0, 0, desc)

    # ── Таблица 3: Перечень закупок ──
    if len(tables) > 3:
        t = tables[3]
        capex_raw = niche["capex_items"].get(format_id, [])
        if not capex_raw:
            # Попробуем первый доступный формат
            for k, v in niche["capex_items"].items():
                capex_raw = v
                break
        items = split_grant_own(capex_raw, grant_amount, own_funds)
        total_grant = 0
        total_own = 0
        total_all = 0
        for i, item in enumerate(items):
            row_idx = i + 1
            if row_idx >= len(t.rows) - 1:
                break
            set_cell(t, row_idx, 0, str(i + 1))
            set_cell(t, row_idx, 1, item["name"])
            set_cell(t, row_idx, 2, str(item["qty"]))
            set_cell(t, row_idx, 3, fmt(item["price"]))
            set_cell(t, row_idx, 4, fmt(item["total"]))
            set_cell(t, row_idx, 5, fmt(item["grant"]) if item["grant"] > 0 else "—")
            set_cell(t, row_idx, 6, fmt(item["own"]) if item["own"] > 0 else "—")
            total_grant += item["grant"]
            total_own += item["own"]
            total_all += item["total"]
        # Итого (последняя строка)
        last_row = len(t.rows) - 1
        set_cell(t, last_row, 1, "ИТОГО")
        set_cell(t, last_row, 4, fmt(total_all))
        set_cell(t, last_row, 5, fmt(total_grant))
        set_cell(t, last_row, 6, fmt(total_own) if total_own > 0 else "—")

    # ── Таблица 4: Цены на услуги ──
    if len(tables) > 4:
        t = tables[4]
        services = niche["services"]
        for i, svc in enumerate(services):
            row_idx = i + 1
            if row_idx >= len(t.rows):
                break
            set_cell(t, row_idx, 0, str(i + 1))
            set_cell(t, row_idx, 1, svc["name"])
            set_cell(t, row_idx, 2, fmt(svc["price"]) + " ₸")

    # ── Таблица 5: Доход в месяц ──
    if len(tables) > 5:
        t = tables[5]
        services = niche["services"]
        monthly_revenue = 0
        for svc in services:
            monthly_revenue += svc["price"] * svc["daily_volume"] * 30
        main_svc = services[0]["name"] if services else "Основная услуга"
        avg_check = sum(s["price"] for s in services) // max(len(services), 1)
        daily_clients = sum(s["daily_volume"] for s in services)
        if len(t.rows) > 1:
            set_cell(t, 1, 0, main_svc)
            set_cell(t, 1, 1, fmt(avg_check) + " ₸")
            set_cell(t, 1, 2, str(daily_clients * 30))
            set_cell(t, 1, 3, fmt(monthly_revenue) + " ₸")
        if len(t.rows) > 2:
            set_cell(t, 2, 0, "Итого")
            set_cell(t, 2, 3, fmt(monthly_revenue) + " ₸")

    # ── Таблица 6: Прогноз прибыли на 3 года ──
    if len(tables) > 6:
        t = tables[6]
        opex = niche["opex_monthly"].get(format_id, {})
        if not opex:
            for k, v in niche["opex_monthly"].items():
                opex = v
                break
        if not opex:
            opex = {"rent": 150000, "fot": 200000, "raw_materials": 150000, "utilities": 25000, "other": 25000}
        monthly_revenue_calc = sum(s["price"] * s["daily_volume"] * 30 for s in niche["services"])
        forecast = calc_profit_forecast(monthly_revenue_calc, opex, start_month)
        # Row 3: доход
        if len(t.rows) > 3:
            set_cell(t, 3, 1, fmt(forecast["y1"]["revenue"]))
            set_cell(t, 3, 2, fmt(forecast["y2"]["revenue"]))
            set_cell(t, 3, 3, fmt(forecast["y3"]["revenue"]))
        # Row 4: итого доход
        if len(t.rows) > 4:
            set_cell(t, 4, 1, fmt(forecast["y1"]["revenue"]))
            set_cell(t, 4, 2, fmt(forecast["y2"]["revenue"]))
            set_cell(t, 4, 3, fmt(forecast["y3"]["revenue"]))
        # Row 6: расход
        if len(t.rows) > 6:
            set_cell(t, 6, 1, fmt(forecast["y1"]["expense"]))
            set_cell(t, 6, 2, fmt(forecast["y2"]["expense"]))
            set_cell(t, 6, 3, fmt(forecast["y3"]["expense"]))
        # Row 7: итого расход
        if len(t.rows) > 7:
            set_cell(t, 7, 1, fmt(forecast["y1"]["expense"]))
            set_cell(t, 7, 2, fmt(forecast["y2"]["expense"]))
            set_cell(t, 7, 3, fmt(forecast["y3"]["expense"]))
        # Row 8: прибыль
        if len(t.rows) > 8:
            set_cell(t, 8, 1, fmt(forecast["y1"]["profit"]))
            set_cell(t, 8, 2, fmt(forecast["y2"]["profit"]))
            set_cell(t, 8, 3, fmt(forecast["y3"]["profit"]))

    # ── Таблица 7: Комментарии ──
    if len(tables) > 7:
        total_opex_m = sum(opex.values()) if opex else 550000
        comment = (
            f"Ежемесячные расходы составляют ~{fmt(total_opex_m)} ₸ "
            f"(аренда, зарплата, сырьё, коммунальные). "
            f"Выручка в первый месяц ожидается ниже плановой (~60-70% от полной мощности). "
            f"Выход на плановый объём — через 2-3 месяца. "
            f"Прибыль реинвестируется в развитие бизнеса и пополнение оборотных средств."
        )
        set_cell(tables[7], 0, 0, comment)

    # ── Таблица 8: Конкуренты ──
    if len(tables) > 8:
        t = tables[8]
        if len(t.rows) > 1:
            set_cell(t, 1, 0, niche["name_ru"])
            avg_price = niche["services"][0]["price"] if niche["services"] else 2000
            set_cell(t, 1, 1, f"от {fmt(avg_price)} ₸")
            set_cell(t, 1, 2, f"Жители города {city['name']}")
            set_cell(t, 1, 3, city["name"])
            set_cell(t, 1, 4, "Качество, скорость, цена")
            set_cell(t, 1, 5, niche["competitors_text"])

    # ── Таблица 9: Анализ потребителей ──
    if len(tables) > 9:
        t = tables[9]
        set_cell(t, 0, 1, fmt(city["avg_wage"]) + " ₸")
        set_cell(t, 1, 1, niche["consumer_annual_purchases"])
        set_cell(t, 2, 1, niche["data_sources"])
        set_cell(t, 3, 1, niche["competitors_text"])

    # ── Таблица 10: Локация ──
    if len(tables) > 10:
        t = tables[10]
        set_cell(t, 0, 1, f"{city['name']}, {location_description}")
        reason = niche["location_reasons"].get(loc_type, "Удобная локация с хорошим трафиком.")
        set_cell(t, 1, 1, reason)
        set_cell(t, 2, 1, niche["channels_text"])

    # ── Таблица 11: Каналы взаимодействия ──
    if len(tables) > 11:
        t = tables[11]
        set_cell(t, 0, 1, niche["suppliers_text"])
        set_cell(t, 1, 1, niche["advantage_text"])

    # ── Сохраняем в bytes ──
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
